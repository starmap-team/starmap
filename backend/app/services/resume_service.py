"""Resume parsing and extraction helpers for phase 4."""

from __future__ import annotations

import hashlib
import json
from io import BytesIO
from typing import Any

import pdfplumber
from docx import Document
from loguru import logger
from redis.asyncio import Redis

from app.core.extraction.jd_extract import extract_from_jd
from app.core.extraction.resume_eval import run_resume_evaluation  # noqa: F401  (re-export:供 api 层经 service 引用)

RESUME_CACHE_PREFIX = "resume-extract:"
RESUME_CACHE_TTL_SECONDS = 24 * 60 * 60
"""LLM 抽取结果内容哈希缓存（24h）。

本地 Ollama 抽取耗时 40-120s+（后端已放宽到 300s），同一简历重复上传时
逐次全量 LLM 抽取体验极差。按文件内容 sha256 缓存 pipeline 结果，命中即秒回。
"""

SUPPORTED_RESUME_EXTENSIONS = {"pdf", "docx"}
"""Whitelist of resume file extensions the backend can parse.

The legacy Microsoft Word ``.doc`` binary format (pre-2007 OLE compound file)
cannot be reliably extracted without third-party tools such as
``antiword``/``wvText``/LibreOffice — none of which are part of the project
runtime. B24: stop accepting ``.doc`` here so callers get a clear 4xx error
instead of garbage output.
"""


def get_resume_extension(filename: str) -> str:
    """Return the lowercased file extension or raise for unsupported names."""
    if not filename or "." not in filename:
        raise ValueError("Resume filename must include an extension")
    return filename.rsplit(".", 1)[-1].lower()


def ensure_supported_resume(filename: str) -> str:
    """Validate the uploaded resume extension."""
    ext = get_resume_extension(filename)
    if ext not in SUPPORTED_RESUME_EXTENSIONS:
        supported = ", ".join(f".{item}" for item in sorted(SUPPORTED_RESUME_EXTENSIONS))
        raise ValueError(f"Unsupported file type: .{ext}. Supported: {supported}")
    return ext


def _decode_text(content_bytes: bytes) -> str:
    for encoding in ("utf-8", "gbk", "utf-16", "latin-1"):
        try:
            text = content_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
        if text.strip():
            return text
    return content_bytes.decode("utf-8", errors="ignore")


def _extract_pdf_text(content_bytes: bytes) -> str:
    pages: list[str] = []
    try:
        with pdfplumber.open(BytesIO(content_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(page_text.strip())
    except (ValueError, TypeError, RuntimeError) as exc:
        logger.warning("PDF parsing failed, fallback to raw decode: {}", exc)
    except Exception as exc:
        logger.exception("Unexpected error in PDF parsing: {}", exc)
    return "\n".join(pages).strip() or _decode_text(content_bytes).strip()


def _extract_docx_text(content_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(content_bytes))
    except (ValueError, TypeError) as exc:
        logger.warning("DOCX parsing failed, fallback to raw decode: {}", exc)
    except Exception as exc:
        logger.exception("Unexpected error in DOCX parsing: {}", exc)
        return _decode_text(content_bytes).strip()

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines).strip()


def extract_resume_text(filename: str, content_bytes: bytes) -> str:
    """Extract plaintext from supported resume formats."""
    ext = ensure_supported_resume(filename)

    if ext == "pdf":
        text = _extract_pdf_text(content_bytes)
    elif ext == "docx":
        text = _extract_docx_text(content_bytes)
    else:
 # Defensive — ensure_supported_resume() already rejects unsupported
 # extensions (including legacy ".doc" — see B24). Reaching here means
 # a future maintainer added an extension to the router without
 # wiring a parser; fall back to plain text decode rather than guess.
        text = _decode_text(content_bytes).strip()

    if not text:
        raise ValueError("Resume file contains no extractable text")
    return text


async def run_resume_extraction(
    filename: str,
    content_bytes: bytes,
    options: dict[str, Any] | None = None,
    redis_client: Redis | None = None,
) -> dict[str, Any]:
    """Parse a resume file and run the shared extraction pipeline.

    Results are cached by file content sha256 when a Redis client is provided
    (the ``/resume/upload`` route injects ``get_redis_client``). Same-file
    re-uploads bypass the slow LLM extraction entirely.
    """
    text = extract_resume_text(filename, content_bytes)
    merged_options = {"source": "resume"}
    if options:
        merged_options.update(options)

    cache_key = _resume_cache_key(content_bytes)
    if redis_client is not None:
        cached = await _read_resume_cache(redis_client, cache_key)
        if cached is not None:
            logger.info("Resume extraction cache HIT ({} bytes) — skipping LLM", len(content_bytes))
            return cached

    result = await extract_from_jd(text, options=merged_options)

    if redis_client is not None and result.get("success"):
        await _write_resume_cache(redis_client, cache_key, result)
    return result


def _resume_cache_key(content_bytes: bytes) -> str:
    digest = hashlib.sha256(content_bytes).hexdigest()
    return f"{RESUME_CACHE_PREFIX}{digest}"


async def _read_resume_cache(redis_client: Redis, key: str) -> dict[str, Any] | None:
    try:
        raw = await redis_client.get(key)
        if not raw:
            return None
        parsed = json.loads(raw)
        return parsed if isinstance(parsed, dict) else None
    except Exception as exc:  # noqa: BLE001 — 缓存失败不阻断主流程
        logger.warning("Resume cache read failed ({}): {}", key, exc)
        return None


async def _write_resume_cache(redis_client: Redis, key: str, result: dict[str, Any]) -> None:
    try:
        await redis_client.set(key, json.dumps(result, ensure_ascii=False), ex=RESUME_CACHE_TTL_SECONDS)
    except Exception as exc:  # noqa: BLE001 — 缓存失败不阻断主流程
        logger.warning("Resume cache write failed ({}): {}", key, exc)
